"""Forager Knowledge MCP Server.

Indexes all project documentation for search and retrieval via Claude Desktop.
"""

from __future__ import annotations

import json
import re
import time
from datetime import date
from pathlib import Path

from mcp.server.fastmcp import FastMCP

from .documents import LoadedDocument, load_all_documents
from .indexer import build_chunks
from .search import KnowledgeSearch


# Resolve repo root (Tools/mcp-knowledge/../../ = repo root)
REPO_ROOT = Path(__file__).resolve().parent.parent.parent.parent

mcp = FastMCP("forager-knowledge")

# Global state — initialized on first tool call
_search_engine: KnowledgeSearch | None = None
_documents: list[LoadedDocument] = []
_index_time: float = 0.0


def _ensure_indexed() -> KnowledgeSearch:
    """Build index on first use (lazy init)."""
    global _search_engine, _documents, _index_time
    if _search_engine is not None:
        return _search_engine

    start = time.time()
    _documents = load_all_documents(REPO_ROOT)
    chunks = build_chunks(_documents)
    _search_engine = KnowledgeSearch(chunks)
    _index_time = time.time() - start
    return _search_engine


VALID_CATEGORIES = [
    "core-doc", "learning-note", "adr", "prd", "newsletter",
    "research", "guideline", "journal", "insight", "testing",
    "test-corpus", "mockup", "project-config", "other",
]


@mcp.tool()
def search_knowledge(
    query: str,
    category: str | None = None,
    max_results: int = 5,
) -> str:
    """Search across all forager project knowledge — docs, learning notes, ADRs, newsletters, journals.

    Args:
        query: Search query (e.g., "CloudKit sharing challenges", "ingredient parsing")
        category: Optional filter — one of: core-doc, learning-note, adr, prd, newsletter, research, guideline, testing, project-config, other
        max_results: Number of results to return (default 5, max 20)
    """
    engine = _ensure_indexed()

    if category and category not in VALID_CATEGORIES:
        return f"Invalid category '{category}'. Valid: {', '.join(VALID_CATEGORIES)}"

    max_results = min(max(1, max_results), 20)
    results = engine.search(query, category=category, max_results=max_results)

    if not results:
        return f"No results found for '{query}'" + (f" in category '{category}'" if category else "")

    output_parts = [f"Found {len(results)} results for '{query}':\n"]
    for i, r in enumerate(results, 1):
        # Truncate excerpt to ~800 chars
        excerpt = r.chunk.text[:800]
        if len(r.chunk.text) > 800:
            excerpt += "..."

        output_parts.append(
            f"--- Result {i} (score: {r.score:.2f}) ---\n"
            f"Title: {r.chunk.display_title}\n"
            f"Category: {r.chunk.doc_category}\n"
            f"File: {r.chunk.doc_path}\n"
            f"\n{excerpt}\n"
        )

    return "\n".join(output_parts)


@mcp.tool()
def read_document(path: str) -> str:
    """Read a full document by path or name substring.

    Args:
        path: Relative path from repo root (e.g., "docs/learning-notes/29-m7-cloudkit-household-journey.md") or a name substring to fuzzy match
    """
    _ensure_indexed()

    # Try exact path match first
    for doc in _documents:
        if doc.file_path == path:
            return _format_document(doc)

    # Try substring match on path or title
    query_lower = path.lower()
    matches = [
        doc for doc in _documents
        if query_lower in doc.file_path.lower() or query_lower in doc.title.lower()
    ]

    if not matches:
        return f"No document found matching '{path}'"

    if len(matches) == 1:
        return _format_document(matches[0])

    # Multiple matches — list them
    lines = [f"Multiple documents match '{path}'. Be more specific:\n"]
    for doc in matches[:15]:
        lines.append(f"  [{doc.category}] {doc.file_path} — {doc.title}")
    return "\n".join(lines)


def _format_document(doc: LoadedDocument) -> str:
    """Format a full document for output."""
    return (
        f"Title: {doc.title}\n"
        f"Category: {doc.category}\n"
        f"File: {doc.file_path}\n"
        f"Size: {doc.size_bytes:,} bytes\n"
        f"{'=' * 60}\n\n"
        f"{doc.content}"
    )


@mcp.tool()
def list_documents(
    category: str | None = None,
    query: str | None = None,
) -> str:
    """List available documents, optionally filtered by category or title.

    Args:
        category: Filter by type — one of: core-doc, learning-note, adr, prd, newsletter, research, guideline, testing, project-config, other
        query: Filter by title substring
    """
    _ensure_indexed()

    if category and category not in VALID_CATEGORIES:
        return f"Invalid category '{category}'. Valid: {', '.join(VALID_CATEGORIES)}"

    docs = _documents
    if category:
        docs = [d for d in docs if d.category == category]
    if query:
        q = query.lower()
        docs = [d for d in docs if q in d.title.lower() or q in d.file_path.lower()]

    if not docs:
        return "No documents found matching filters."

    # Group by category
    by_category: dict[str, list[LoadedDocument]] = {}
    for doc in docs:
        by_category.setdefault(doc.category, []).append(doc)

    lines = [f"Found {len(docs)} documents:\n"]
    for cat in sorted(by_category.keys()):
        lines.append(f"\n[{cat}] ({len(by_category[cat])} files)")
        for doc in by_category[cat]:
            size_kb = doc.size_bytes / 1024
            lines.append(f"  {doc.file_path} — {doc.title} ({size_kb:.0f} KB)")

    lines.append(f"\nIndex: {len(_documents)} docs, {len(_search_engine.chunks) if _search_engine else 0} chunks, built in {_index_time:.2f}s")
    return "\n".join(lines)


@mcp.tool()
def get_project_status() -> str:
    """Get current project status — active milestone, branch, and next priorities."""
    _ensure_indexed()

    parts = []
    for doc in _documents:
        if doc.file_path == "docs/current-story.md":
            # First ~150 lines
            lines = doc.content.split("\n")[:150]
            parts.append("=== CURRENT STORY ===\n" + "\n".join(lines))
        elif doc.file_path == "docs/next-prompt.md":
            lines = doc.content.split("\n")[:80]
            parts.append("\n=== NEXT PROMPT ===\n" + "\n".join(lines))

    if not parts:
        return "Could not find current-story.md or next-prompt.md"

    return "\n".join(parts)


@mcp.tool()
def get_capabilities() -> str:
    """List all OpenSpec capabilities with one-line summaries.

    Scans openspec/specs/<capability>/spec.md. Returns each capability's
    kebab-case name, path, and a summary extracted from the spec's Overview
    or first requirement. Use this to orient to what the system does before
    proposing or implementing changes.
    """
    _ensure_indexed()

    specs_dir = REPO_ROOT / "openspec" / "specs"
    if not specs_dir.exists():
        return "No openspec/specs/ directory found"

    entries = []
    for cap_dir in sorted(specs_dir.iterdir()):
        if not cap_dir.is_dir():
            continue
        spec_path = cap_dir / "spec.md"
        if not spec_path.exists():
            continue

        name = cap_dir.name
        rel_path = f"openspec/specs/{name}/spec.md"
        summary = _extract_capability_summary(spec_path)
        entries.append((name, rel_path, summary))

    if not entries:
        return "No capabilities found"

    lines = [f"Found {len(entries)} capabilities:\n"]
    for name, path, summary in entries:
        lines.append(f"  [{name}] {path}")
        lines.append(f"    {summary}")
    return "\n".join(lines)


def _extract_capability_summary(spec_path: Path) -> str:
    """Pull a one-line summary from a spec file — Overview first line, or first Requirement description."""
    try:
        content = spec_path.read_text(encoding="utf-8")
    except Exception as e:
        return f"(unreadable: {e})"

    # Try Overview section first
    overview_match = re.search(r"^##\s+Overview\s*\n+(.+?)(?=\n\n|\n##|\Z)", content, re.MULTILINE | re.DOTALL)
    if overview_match:
        first_line = overview_match.group(1).strip().split("\n")[0]
        return first_line[:250]

    # Fall back to first Requirement description
    req_match = re.search(r"^###\s+Requirement:\s*(.+?)\n+(.+?)(?=\n\n|\n####|\n###|\Z)", content, re.MULTILINE | re.DOTALL)
    if req_match:
        req_title = req_match.group(1).strip()
        req_desc = req_match.group(2).strip().split("\n")[0]
        return f"{req_title} — {req_desc[:200]}"

    return "(no Overview or Requirement found)"


@mcp.tool()
def get_services() -> str:
    """List top-level Swift services in Services/ with role hints.

    Scans Services/*.swift (excluding subdirectories like Parsing/, Import/,
    Persistence/). For each file, extracts a role hint from the top-level
    doc comment or class declaration. Useful for orienting to the service
    layer before proposing new services (run /service-check first).
    """
    services_dir = REPO_ROOT / "Services"
    if not services_dir.exists():
        return "No Services/ directory found"

    entries = []
    for service_file in sorted(services_dir.glob("*.swift")):
        name = service_file.name
        rel_path = f"Services/{name}"
        role = _extract_service_role(service_file)
        entries.append((name, rel_path, role))

    if not entries:
        return "No services found"

    lines = [f"Found {len(entries)} top-level services:\n"]
    for name, path, role in entries:
        lines.append(f"  [{name.replace('.swift', '')}] {path}")
        lines.append(f"    {role}")
    return "\n".join(lines)


def _extract_service_role(service_path: Path) -> str:
    """Extract role hint from a Swift service file — top doc comment or class declaration line."""
    try:
        content = service_path.read_text(encoding="utf-8")
    except Exception as e:
        return f"(unreadable: {e})"

    # Try top-level doc comment block (/// lines)
    doc_lines = []
    for line in content.split("\n"):
        stripped = line.strip()
        if stripped.startswith("///"):
            doc_lines.append(stripped.lstrip("/").strip())
        elif stripped.startswith("//"):
            # Plain comments, include up to first break
            doc_lines.append(stripped.lstrip("/").strip())
        elif doc_lines and stripped and not stripped.startswith("import"):
            # Hit non-comment non-import; stop
            break

    if doc_lines:
        # First non-empty doc line
        for d in doc_lines:
            if d:
                return d[:250]

    # Fall back to first class/struct declaration
    class_match = re.search(r"^(?:final\s+)?(?:class|struct|actor|enum)\s+(\w+)", content, re.MULTILINE)
    if class_match:
        return f"declares {class_match.group(1)}"

    # Fall back to file name
    return service_path.stem


@mcp.tool()
def get_adr(number: str) -> str:
    """Fetch a specific ADR by number prefix.

    Scans docs/architecture/ for a file matching the given number prefix
    (e.g., "013" matches "013-scope-aware-fetch-pattern.md"). Returns the
    ADR's number, title, status, path, and full content.

    Args:
        number: The ADR number as a zero-padded string (e.g., "007", "013")
    """
    adr_dir = REPO_ROOT / "docs" / "architecture"
    if not adr_dir.exists():
        return "No docs/architecture/ directory found"

    # Normalize — accept "7", "07", or "007"
    normalized = number.strip().zfill(3)

    for adr_file in sorted(adr_dir.glob(f"{normalized}-*.md")):
        return _format_adr(adr_file)

    # Also try unpadded, since user might pass "7"
    for adr_file in sorted(adr_dir.glob(f"{number.strip()}-*.md")):
        return _format_adr(adr_file)

    available = [f.stem for f in sorted(adr_dir.glob("*.md")) if re.match(r"^\d+", f.stem)]
    return f"No ADR found matching number '{number}'. Available: {', '.join(available)}"


def _format_adr(adr_file: Path) -> str:
    """Format an ADR file for output with extracted metadata."""
    try:
        content = adr_file.read_text(encoding="utf-8")
    except Exception as e:
        return f"Could not read {adr_file}: {e}"

    # Extract title from first heading
    title_match = re.search(r"^#\s+(.+)$", content, re.MULTILINE)
    title = title_match.group(1).strip() if title_match else adr_file.stem

    # Extract status if declared (e.g., **Status**: SUPERSEDED or Status: Active)
    status_match = re.search(r"(?:^|\n)\*\*Status\*\*:\s*([^\n]+)", content)
    if not status_match:
        status_match = re.search(r"(?:^|\n)Status:\s*([^\n]+)", content)
    status = status_match.group(1).strip() if status_match else "Active (no explicit status declared)"

    # Extract number
    number_match = re.match(r"^(\d+)-", adr_file.stem)
    number = number_match.group(1) if number_match else "?"

    rel_path = f"docs/architecture/{adr_file.name}"
    return (
        f"ADR {number}: {title}\n"
        f"Status: {status}\n"
        f"File: {rel_path}\n"
        f"{'=' * 60}\n\n"
        f"{content}"
    )


@mcp.tool()
def get_active_work() -> str:
    """Return parsed extract of current active work — milestone, branch, next action.

    Reads docs/current-story.md (header + active-milestone section) and
    docs/next-prompt.md (active section). Use to quickly orient to what
    is being worked on right now without reading the full files.
    """
    _ensure_indexed()

    parts = []

    # current-story.md — parse header lines + first ACTIVE section
    current_story = next((d for d in _documents if d.file_path == "docs/current-story.md"), None)
    if current_story:
        parts.append("=== ACTIVE WORK (from current-story.md) ===\n")
        lines = current_story.content.split("\n")
        # Header block — first ~10 lines with status / branch / launch path
        parts.extend(lines[:15])
        # First ACTIVE: section if present
        in_active = False
        active_lines: list[str] = []
        for line in lines[15:]:
            if re.match(r"^##\s+ACTIVE:", line):
                in_active = True
                active_lines.append(line)
            elif in_active:
                if re.match(r"^##\s+", line):
                    break
                active_lines.append(line)
        if active_lines:
            parts.append("")
            parts.extend(active_lines[:40])

    # next-prompt.md — first 60 lines (header + Active + first backlog entries)
    next_prompt = next((d for d in _documents if d.file_path == "docs/next-prompt.md"), None)
    if next_prompt:
        parts.append("\n\n=== NEXT PROMPT (from next-prompt.md) ===\n")
        parts.extend(next_prompt.content.split("\n")[:60])

    if not parts:
        return "Could not find current-story.md or next-prompt.md in index"

    return "\n".join(parts)


@mcp.tool()
def get_newsletter_context(
    topic: str,
    include_previous: bool = True,
) -> str:
    """Pull relevant project knowledge for newsletter writing.

    Searches project docs for the topic and optionally includes excerpts from
    previous newsletters for style reference.

    Args:
        topic: What the newsletter section is about (e.g., "CloudKit sharing", "ML parsing")
        include_previous: Include excerpts from recent newsletters for tone/style (default true)
    """
    engine = _ensure_indexed()

    parts = ["=== PROJECT CONTEXT FOR NEWSLETTER ===\n"]
    parts.append(f"Topic: {topic}\n")

    # Search project docs (excluding newsletters)
    all_results = engine.search(topic, max_results=15)
    project_results = [r for r in all_results if r.chunk.doc_category != "newsletter"][:8]

    if project_results:
        parts.append("\n--- Relevant Project Knowledge ---\n")
        for r in project_results:
            excerpt = r.chunk.text[:600]
            if len(r.chunk.text) > 600:
                excerpt += "..."
            parts.append(
                f"[{r.chunk.doc_category}] {r.chunk.display_title}\n"
                f"File: {r.chunk.doc_path}\n"
                f"{excerpt}\n"
            )

    # Search development journal
    journal_results = engine.search(topic, category="core-doc", max_results=3)
    journal_hits = [r for r in journal_results if "journal" in r.chunk.doc_path.lower()]
    if journal_hits:
        parts.append("\n--- Development Journal Entries ---\n")
        for r in journal_hits:
            excerpt = r.chunk.text[:500]
            if len(r.chunk.text) > 500:
                excerpt += "..."
            parts.append(f"{r.chunk.display_title}\n{excerpt}\n")

    # Include recent newsletter excerpts for style
    if include_previous:
        newsletter_docs = [d for d in _documents if d.category == "newsletter"]
        # Take last 2-3 newsletters
        recent = newsletter_docs[-3:] if len(newsletter_docs) >= 3 else newsletter_docs
        if recent:
            parts.append("\n--- Recent Newsletter Style Reference ---\n")
            for doc in recent:
                # First ~500 chars for style reference
                excerpt = doc.content[:500]
                if len(doc.content) > 500:
                    excerpt += "..."
                parts.append(f"[{doc.title}]\n{excerpt}\n")

    return "\n".join(parts)


@mcp.tool()
def draft_newsletter_section(
    topic: str,
    style_notes: str | None = None,
    max_words: int = 500,
) -> str:
    """Get a curated context bundle and outline for drafting a newsletter section.

    Returns relevant facts, quotes, and a suggested structure. Use this context
    to write the actual prose in the conversation.

    Args:
        topic: Section topic (e.g., "the challenge of CloudKit sync debugging")
        style_notes: Optional style guidance (e.g., "more technical", "focus on the human story")
        max_words: Target word count for the section (default 500)
    """
    engine = _ensure_indexed()

    parts = [f"=== NEWSLETTER DRAFT CONTEXT ===\n"]
    parts.append(f"Topic: {topic}")
    parts.append(f"Target: ~{max_words} words")
    if style_notes:
        parts.append(f"Style: {style_notes}")
    parts.append("")

    # Gather relevant project knowledge
    results = engine.search(topic, max_results=10)
    project_results = [r for r in results if r.chunk.doc_category != "newsletter"]

    if project_results:
        parts.append("--- Key Facts & Context ---\n")
        for r in project_results[:6]:
            excerpt = r.chunk.text[:500]
            if len(r.chunk.text) > 500:
                excerpt += "..."
            parts.append(
                f"Source: [{r.chunk.doc_category}] {r.chunk.display_title}\n"
                f"{excerpt}\n"
            )

    # Check for relevant insights
    insight_results = engine.search(topic, category="core-doc", max_results=3)
    insight_hits = [r for r in insight_results if "insight" in r.chunk.doc_path.lower()]
    if insight_hits:
        parts.append("--- Insights Log Entries ---\n")
        for r in insight_hits:
            parts.append(f"{r.chunk.text[:400]}\n")

    # Suggest structure
    parts.append("--- Suggested Structure ---\n")
    parts.append(f"1. Hook — Open with the specific challenge or moment related to '{topic}'")
    parts.append(f"2. Context — What was being built and why it mattered")
    parts.append(f"3. The Turn — What was discovered, decided, or changed")
    parts.append(f"4. Takeaway — The lesson or principle that generalizes beyond forager")
    parts.append(f"\nUse the context above to write ~{max_words} words in the newsletter's established voice.")

    return "\n".join(parts)


@mcp.tool()
def create_newsletter_draft(
    title: str,
    content: str,
    number: int | None = None,
    output_dir: str | None = None,
) -> str:
    """Generate a .docx newsletter file from provided content.

    Args:
        title: Newsletter title (used for filename and document heading)
        content: Newsletter body (markdown-formatted — headings, bold, italic, lists converted to docx)
        number: Issue number (auto-increments from existing newsletters if omitted)
        output_dir: Where to save (default: docs/newsletters/)
    """
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    target_dir = Path(output_dir) if output_dir else REPO_ROOT / "docs" / "newsletters"
    target_dir.mkdir(parents=True, exist_ok=True)

    # Auto-detect next issue number
    if number is None:
        existing = sorted(target_dir.glob("*.docx"))
        max_num = 0
        for f in existing:
            match = re.search(r"-\s*(\d+)\s*-", f.name)
            if match:
                max_num = max(max_num, int(match.group(1)))
        number = max_num + 1

    # Build filename
    today = date.today().strftime("%Y.%m.%d")
    filename = f"{today} - {number:03d} - {title}.docx"
    filepath = target_dir / filename

    # Create document
    doc = DocxDocument()

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Convert markdown content to docx
    _markdown_to_docx(doc, content)

    doc.save(str(filepath))

    return f"Newsletter draft created: {filepath.relative_to(REPO_ROOT)}\nAbsolute path: {filepath}"


def _markdown_to_docx(doc: DocxDocument, content: str) -> None:
    """Convert markdown-formatted text to docx paragraphs."""
    from docx.shared import Pt

    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        # Unordered list
        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            text = re.sub(r"^[\s]*[-*]\s+", "", line)
            para = doc.add_paragraph(style="List Bullet")
            _add_inline_formatting(para, text)
        # Numbered list
        elif re.match(r"^\s*\d+\.\s+", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            para = doc.add_paragraph(style="List Number")
            _add_inline_formatting(para, text)
        # Blank line
        elif not line.strip():
            pass  # skip blank lines (paragraph spacing handles this)
        # Regular paragraph
        else:
            para = doc.add_paragraph()
            _add_inline_formatting(para, line)

        i += 1


def _add_inline_formatting(paragraph, text: str) -> None:
    """Handle **bold** and *italic* inline formatting."""
    # Split on bold and italic markers
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def main():
    """Entry point for the MCP server."""
    mcp.run()


if __name__ == "__main__":
    main()
