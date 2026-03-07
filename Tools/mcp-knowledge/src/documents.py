"""Document loading for markdown and docx files."""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document as DocxDocument


CATEGORY_RULES: list[tuple[str, str]] = [
    ("docs/newsletters/", "newsletter"),
    ("docs/learning-notes/", "learning-note"),
    ("docs/architecture/", "adr"),
    ("docs/prds/", "prd"),
    ("docs/import-research/", "research"),
    ("docs/ux-research/", "research"),
    ("docs/test-corpus/", "test-corpus"),
    ("docs/testing/", "testing"),
    ("docs/mockups/", "mockup"),
    ("CLAUDE.md", "project-config"),
]

CORE_DOCS = {
    "current-story.md",
    "roadmap.md",
    "requirements.md",
    "project-index.md",
    "insights-log.md",
    "development-journal.md",
    "next-prompt.md",
}

GUIDELINE_DOCS = {
    "development-guidelines.md",
    "session-startup-checklist.md",
    "project-naming-standards.md",
    "git-workflow-for-milestones.md",
    "git-cheatsheet.md",
}


@dataclass
class LoadedDocument:
    """A document loaded from disk with metadata."""

    title: str
    category: str
    file_path: str  # relative to repo root
    content: str
    size_bytes: int


def classify_document(rel_path: str) -> str:
    """Determine category from relative path."""
    for prefix, category in CATEGORY_RULES:
        if rel_path.startswith(prefix):
            return category

    filename = Path(rel_path).name
    if filename in CORE_DOCS:
        return "core-doc"
    if filename in GUIDELINE_DOCS:
        return "guideline"
    return "other"


def extract_title_from_markdown(content: str, filename: str) -> str:
    """Extract title from first H1 heading or fall back to filename."""
    for line in content.split("\n")[:10]:
        line = line.strip()
        if line.startswith("# ") and not line.startswith("## "):
            return line[2:].strip()
    return Path(filename).stem


def extract_title_from_docx(filename: str) -> str:
    """Extract title from docx filename convention: YYYY.MM.DD - NNN - Title.docx"""
    stem = Path(filename).stem
    # Match pattern: date - number - title
    match = re.match(r"\d{4}\.\d{2}\.\d{2}\s*-\s*\d+\s*-\s*(.+)", stem)
    if match:
        return match.group(1).strip()
    return stem


def load_markdown(path: Path, repo_root: Path) -> LoadedDocument:
    """Load a markdown file."""
    content = path.read_text(encoding="utf-8", errors="replace")
    rel_path = str(path.relative_to(repo_root))
    return LoadedDocument(
        title=extract_title_from_markdown(content, path.name),
        category=classify_document(rel_path),
        file_path=rel_path,
        content=content,
        size_bytes=path.stat().st_size,
    )


def load_docx(path: Path, repo_root: Path) -> LoadedDocument:
    """Load a .docx file, extracting all paragraph text."""
    doc = DocxDocument(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    content = "\n\n".join(paragraphs)
    rel_path = str(path.relative_to(repo_root))
    return LoadedDocument(
        title=extract_title_from_docx(path.name),
        category=classify_document(rel_path),
        file_path=rel_path,
        content=content,
        size_bytes=path.stat().st_size,
    )


def load_all_documents(repo_root: Path) -> list[LoadedDocument]:
    """Load all markdown and docx documents from the repo."""
    docs: list[LoadedDocument] = []

    # Load CLAUDE.md from repo root
    claude_md = repo_root / "CLAUDE.md"
    if claude_md.exists():
        docs.append(load_markdown(claude_md, repo_root))

    # Load all markdown in docs/
    docs_dir = repo_root / "docs"
    if docs_dir.exists():
        for md_path in sorted(docs_dir.rglob("*.md")):
            try:
                docs.append(load_markdown(md_path, repo_root))
            except Exception:
                continue

    # Load all docx in docs/newsletters/
    newsletters_dir = repo_root / "docs" / "newsletters"
    if newsletters_dir.exists():
        for docx_path in sorted(newsletters_dir.rglob("*.docx")):
            try:
                docs.append(load_docx(docx_path, repo_root))
            except Exception:
                continue

    return docs
